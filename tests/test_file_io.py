import os
import tempfile
from docx import Document


def test_save_txt_writes_content():
    from file_io import save_txt
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as f:
        path = f.name
    try:
        save_txt("hello world", path)
        with open(path, encoding="utf-8") as f:
            assert f.read() == "hello world"
    finally:
        os.unlink(path)


def test_save_txt_overwrites_existing():
    from file_io import save_txt
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as f:
        f.write("old content")
        path = f.name
    try:
        save_txt("new content", path)
        with open(path, encoding="utf-8") as f:
            assert f.read() == "new content"
    finally:
        os.unlink(path)


def test_save_docx_writes_paragraphs():
    from file_io import save_docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        save_docx("line one\nline two", path)
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        assert "line one" in paragraphs
        assert "line two" in paragraphs
    finally:
        os.unlink(path)


def test_save_docx_empty_text():
    from file_io import save_docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        save_docx("", path)
        doc = Document(path)
        # Empty string produces one empty paragraph — no crash
        assert isinstance(doc.paragraphs, list)
    finally:
        os.unlink(path)
